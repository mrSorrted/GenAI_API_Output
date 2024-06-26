#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Jun 25 23:58:07 2024


Final code by importing data from JSON

@author: anu
"""

# Has functions to work with json dat
import json
# Has functions to worl with files/directories
import os
# Helps to create Word Document
from docx import Document
# Helps to work on OOXML elements
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Function to set the width of each column in the table
def set_column_widths(table):
    # A list to store each column width
    column_widths = []
    # To iterate all the columns in a table
    for col_index, column in enumerate(table.columns):
        # To get the Maximum Width of the cell present in the column
        max_width = max(cell.width for cell in column.cells)
        # Add the max width into the list
        column_widths.append(max_width)
    
    # Ṭo iterate all the columns in a table
    for col_index, column in enumerate(table.columns):
        for cell in column.cells:
            # To set width od each cell with the max width cell present in the column
            cell.width = column_widths[col_index]


# Function to add background color to the row that is needed
def set_row_bg_color(row, color):
    # To iterate each cell for a given row
    for cell in row.cells:
        # To get existing cell property or create a new cell property
        tablecellprop = cell._element.get_or_add_tcPr()
        # To create xml for adding background color 
        bg = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
        # Set the Background color to the cell
        tablecellprop.append(bg)

# Function to add background color to the column taht is needed
def set_column_bg_color(table, col_index, color):
    # To iterate each cell for a given row
    for row in table.rows:
        # To get the index of the column that needs to be set with background color
        cell = row.cells[col_index]
        # To get existing cell property or create a new cell property
        tablecellprop = cell._element.get_or_add_tcPr()
        # To create xml for adding background color
        bg = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
        # Set the Background color to the cell
        tablecellprop.append(bg)

 # Function to create a table from the json file and print the table to a Word doc   
def create_table_from_jsonfile(json_data, row1_headings, row2_headings, row_names, doc):
    # To get Total Number of Rows (2 Heading Rows + Actual Data Rows)
    num_of_rows = len(row_names) + 2
    # To get Total Number od Columns (1 Heading Column + Actual Data Columns)
    num_of_cols = len(row1_headings) + 1
    
    # To add table into Word doc with obtained number of rows & columns
    # Style as table grid - Adds Grid layout to table
    table = doc.add_table(rows=num_of_rows, cols=num_of_cols, style='Table Grid')

    # To set Top Left cell heading
    table.cell(0, 0).text = "USD m"
    
    # To set the headings for the First Row
    for i, heading in enumerate(row1_headings):
        table.cell(0, i + 1).text = heading

    # To set Second Top Left cell heading
    table.cell(1, 0).text = "USD"
    
    # To set the headings for the Second Row
    for i, heading in enumerate(row2_headings):
        table.cell(1, i + 1).text = heading

    # Fetches the Row Name from JSON and populate in First Column
    # Get the Row's Data from JSON and populates remaining columns
    for i, row_name in enumerate(row_names):
        table.cell(i + 2, 0).text = row_name
        row_data = json_data.get(row_name, [])
        for j, value in enumerate(row_data):
            # To not get Index out of range error
            if j < len(row1_headings):
                table.cell(i + 2, j + 1).text = str(value) if value is not None else ""
                                         
    # To set background color for specific rows
    # Specify the Row Indexes where you want the background color
    for row_index in [0, 1, 11, 18, 25, 30]:  
        set_row_bg_color(table.rows[row_index], '000080')  # Navy blue

    # Set background color for specific columns
    # Specify the Column Indexes where you want the background color
    for col_index in [0]:
        set_column_bg_color(table, col_index, '99CCFF')  # Light Blue
        
    set_column_widths(table)
        
    
# Read the JSON data from the JSON file
with open('/Users/anu/.spyder-py3/data2.json') as f:
    json_data = json.load(f)

# Predefined Row/Column Names for the table
row1_headings = ['FY19', 'FY20', 'FY21', 'FY22', 'FY23', 'FY24e', 'FY25e', 'FY26e']
row2_headings = ['Audited', 'Audited', 'Audited', 'Audited', 'Audited', 'HSBC Research', ' ', ' ']
row_names = ['Sales Revenue', 'Gross Profit', 'Gross Profit Margin %', 'EBITDA', 
             'EBITDA Margin %', 'Operating Profit', 'Operating Profit Margin %', 
             'Interest Expense (Net)', 'Net Profit', 'Balance Sheet', 'Net Worth',
             'Tangible Net Worth', 'Total Ext. Funded Debt (TFD)', 'Cash + Mkt. Securities',
             'Net Debt', 'Working Capital Days', 'Cash Flow', 'Change in Working Capital',
             'Operating Cash Flow', 'Investing Cash Flow', 'Financing Cash Flow', 'w/w Dividends',
             'Free Cash Flow', 'Ratios', 'Ext. Gearing (TFD/TNW) (x)', 'Net Debt/EBITDA (x)',
             'TFD/EBITDA (x)', 'NOCF/Interest (x)', 'Financial Covenants', 'TFD/EBITDA-Max[x]',
             'DSCR-Min[x]'
             ]

# To create a new Word Document
doc = Document()

# To give the title in the document
doc.add_heading('Financial Analysis Doc', level=1)

# To create table from JSON data
create_table_from_jsonfile(json_data, row1_headings, row2_headings, row_names, doc)

# Get the current working directory
current_working_directory = os.getcwd()

# Specify the path where the document need to be saved
save_path = os.path.join(current_working_directory, 'Financial Analysis Doc.docx')

# Save the document
doc.save(save_path)

#print("Document created successfully!!")
